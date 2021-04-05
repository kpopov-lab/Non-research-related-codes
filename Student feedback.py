import getpass
import pandas as pd
import numpy as np
import docx
import os
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

username = getpass.getuser()
folder_name = 'Student feedback GE2 2021'
desktop = f'C:/Users/{username}/Desktop'
cloud = f'C:/Users/{username}/OneDrive - University of Edinburgh'


def make_folder():
    # Function to create a folder if it doesn't exist
    try:
        os.makedirs(f'{cloud}/{folder_name}')
    except FileExistsError:
        pass


# Dictionary with the grades for each mark
grades = {range(40): 'F', range(40, 50): 'D', range(50, 60): 'C', range(60, 70): 'B',
          range(70, 80): 'A3', range(80, 90): 'A2', range(90, 101): 'A1'}


class Feedback:
    # Class to create feedback comments and assign a mark and grade

    def __init__(self, last_name, first_name, topic):
        self.last_name = last_name
        self.first_name = first_name
        self.full_name = self.first_name + ' ' + self.last_name
        self.topic = topic
        self.results = {'mark': 0, 'grade': '', 'comments': ''}

    def feedback_word(self):
        # Making comments and compiling them in word format

        self.results['comments'] = input('Insert feedback comments here:\n')
        self.results['mark'] = int(input('What mark do you want to assign? '))  # Assign mark

        # Assign grade based on the mark given above
        for key in grades.keys():
            if self.results['mark'] in key:
                self.results['grade'] = grades[key]
                break

        make_folder()  # Make the folder if it doesn't exist

        # Get list of file names from the target folder
        files = os.listdir(f'{cloud}/{folder_name}')

        # Create feedback file if it doesn't already exist and format it
        if 'Feedback.docx' not in files:
            feedback = docx.Document()
            p = feedback.add_paragraph()
            formatting = p.paragraph_format
            formatting.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            runner = p.add_run('Geotechnical Engineering 2 long report feedback (2021)')
            runner.bold = True
            font = runner.font
            font.name = 'Cambria'
            font.size = Pt(22)
            runner = p.add_run('\nCourse organiser: Prof Jin Sun\nUniversity of Edinburgh\n')
            runner.italic = True
            font = runner.font
            font.name = 'Cambria'
            font.size = Pt(16)

        # Open feedback file if already exists
        else:
            feedback = docx.Document(f'{cloud}/{folder_name}/Feedback.docx')

        # Generate feedback comments and format them
        p = feedback.add_paragraph()

        p.add_run(f'Student name: ')
        runner = p.add_run(f'{self.last_name}, {self.first_name}')
        runner.bold = True
        runner.italic = True
        p.add_run(f"\nTopic: {self.topic}\nStudent mark: {self.results['mark']}%\n"
                  f"Student grade: {self.results['grade']}\n\nComments: {self.results['comments']}\n"
                  f"--------------------------------------------")

        # Check if the feedback is open
        try:
            feedback.save(f'{cloud}/{folder_name}/Feedback.docx')
        except PermissionError:
            print('The file you are trying to edit is currently open and changes cannot be saved.'
                  'Please close file and restart the operation.')

        return self.results['mark'], self.results['grade']


class UpdateSpreadsheet(Feedback):
    # Class to add student mark and grade to a spreadsheet and estimate the average mark

    def __init__(self, last_name, first_name, topic):
        super().__init__(last_name, first_name, topic)
        self.data = pd.read_csv(f'{cloud}/{folder_name}/GE2 marks.csv')

    def update_marks(self):
        # Prepare new data to add to existing file
        print(self.results['mark'])
        new_mark = pd.DataFrame([self.last_name, self.first_name, self.topic, mark, grade],
                                index=self.data.columns).transpose()

        # Add the data without the existing average value
        self.data = pd.concat([self.data.iloc[:-1], new_mark])

        # Estimate the new average and add it to the data
        avg = np.around(self.data.Mark.mean(), 2)
        avg = pd.DataFrame(['N/A', 'N/A', 'Average', avg, 'N/A'], index=self.data.columns).transpose()
        self.data = pd.concat([self.data, avg])

        # Export the new spreadsheet
        self.data.to_csv(f'{cloud}/{folder_name}/GE2 marks.csv', index=False)


a = True

while a:
    l_name = input('Enter last name here: ')
    f_name = input('Enter first name here: ')
    student_topic = input('Enter the topic the student wrote about: ')

    StudentFeedback = Feedback(l_name, f_name, student_topic)
    mark, grade = StudentFeedback.feedback_word()
    UpdateSpreadsheet(l_name, f_name, student_topic).update_marks()

